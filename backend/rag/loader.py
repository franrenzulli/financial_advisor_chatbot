import os
from pathlib import Path
from typing import List
import pytesseract

#from ibm_docprocessing import DocumentParser
from docx import Document as DocxDocument
from PIL import Image
from langchain_core.documents import Document

from rag.splitter import split_documents

PDF_DIR = Path(__file__).resolve().parent.parent / "data"
TXT_DIR = Path(__file__).resolve().parent.parent / "data" / "generated_texts"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png"}


def extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()

    #if ext == ".pdf":
        #parser = DocumentParser()
        #result = parser.parse(file_path)
        #return result.text

    if ext in {".docx", ".doc"}:
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    elif ext in {".jpg", ".jpeg", ".png"}:
        image = Image.open(file_path)
        return pytesseract.image_to_string(image)

    else:
        raise ValueError(f"File type not supported: {ext}")


def load_and_split_from_folder(chunk_size=1000, chunk_overlap=100) -> List[Document]:
    TXT_DIR.mkdir(parents=True, exist_ok=True)
    all_chunks = []

    for file_path in PDF_DIR.iterdir():
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            print(f"Ignored (extension not supported): {file_path.name}")
            continue

        print(f"Processing: {file_path.name}")
        try:
            text = extract_text(file_path)

            # Save the text for reference 
            output_path = TXT_DIR / f"{file_path.stem}.txt"
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"✅ Saved plain text at: {output_path.name}")

            # Create a Document with metadata
            doc = Document(
                page_content=text,
                metadata={"source": str(file_path.name)}
            )

            # Aply the splitter
            chunks = split_documents([doc], chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            all_chunks.extend(chunks)

        except Exception as e:
            print(f"❌ Error processing {file_path.name}: {e}")

    return all_chunks


if __name__ == "__main__":
    chunks = load_and_split_from_folder()
    print(f"\n✅ Total chunks generated: {len(chunks)}")
